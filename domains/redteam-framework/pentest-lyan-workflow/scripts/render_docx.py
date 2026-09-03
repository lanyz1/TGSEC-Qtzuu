#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""渗透测试报告 md -> docx 渲染器 v2。"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 2:
    sys.exit("用法: python3 render_docx_v2.py <report.md> [output.docx]")
SRC = sys.argv[1]
if not SRC.endswith(".md"):
    sys.exit("输入必须是 .md 文件")
OUT = sys.argv[2] if len(sys.argv) > 2 else SRC[:-3] + ".docx"

# ── 颜色系统 ──────────────────────────────────────────────
INK      = "1F3A68"   # 深蓝：正文标题
ACCENT   = "2E6FD4"   # 亮蓝：次级标题竖条
MUTED    = "595959"   # 灰：封面 meta
HDR_BG   = "D9E2F3"  # 表头浅蓝
CODE_BG  = "F4F4F4"  # 代码块灰底
HAIR     = "BFBFBF"  # 分割线
PAGE_BG  = "FAFBFE"  # 封面底色（轻微）

SEVERITY_COLORS = {
    "严重": ("C0392B", "FDECEA"),  # 红
    "高危": ("D35400", "FEF0E7"),  # 橙
    "中危": ("B7770D", "FEF9E7"),  # 黄
    "低危": ("1A7A4A", "E9F7EF"),  # 绿
    "信息": ("2471A3", "EBF5FB"),  # 蓝
}

CJK = "微软雅黑"; LATIN = "Segoe UI"; MONO = "Consolas"

# ── 字体工具 ──────────────────────────────────────────────
def set_font(run, name=LATIN, cjk=CJK, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); r.append(rFonts)
    rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name); rFonts.set(qn('w:eastAsia'), cjk)
    if size  is not None: run.font.size = Pt(size)
    if bold  is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)

def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    old = pPr.find(qn('w:shd'))
    if old is not None: pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:shd'))
    if old is not None: tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def left_bar(paragraph, color=INK, sz='24', space='8'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), sz)
    lf.set(qn('w:space'), space); lf.set(qn('w:color'), color)
    pBdr.append(lf); pPr.append(pBdr)

def set_cell_border(cell, color="D0D7E2", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), r_id)
    run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '2471A3'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    run.append(rPr); t = OxmlElement('w:t'); t.text = text
    t.set(qn('xml:space'), 'preserve'); run.append(t)
    h.append(run); paragraph._p.append(h)

HTML_TAG_RE = re.compile(r'<[^>]+>')
def strip_html(text): return HTML_TAG_RE.sub('', text)

INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|~~[^~]+~~)')

def add_inline(paragraph, text, size=10.5, color=None, bold=False):
    text = strip_html(text); pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_font(r, size=size, color=color, bold=bold if bold else None)
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); set_font(r, size=size, bold=True, color=color)
        elif tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1]); set_font(r, name=MONO, cjk=MONO, size=size-0.5, color='C0392B')
        elif tok.startswith('~~'):
            r = paragraph.add_run(tok[2:-2]); set_font(r, size=size, color=color); r.font.strike = True
        elif tok.startswith('['):
            mm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            if mm: add_hyperlink(paragraph, mm.group(1), mm.group(2))
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_font(r, size=size, color=color, bold=bold if bold else None)

def add_code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15; p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.2)
        shade(p, CODE_BG)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        lf = OxmlElement('w:left')
        lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '16')
        lf.set(qn('w:space'), '6'); lf.set(qn('w:color'), ACCENT)
        pBdr.append(lf); pPr.append(pBdr)
        r = p.add_run(ln if ln != "" else " ")
        set_font(r, name=MONO, cjk=MONO, size=9, color='2C3E50')

def flush_table(doc, rows):
    if not rows: return
    parsed = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    body = [p for p in parsed if not all(re.match(r'^:?-+:?$', c or '') for c in p)]
    if not body: return
    ncols = max(len(r) for r in body)
    tbl = doc.add_table(rows=len(body), cols=ncols)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; tbl.autofit = True
    for ri, row in enumerate(body):
        for ci in range(ncols):
            cell = tbl.cell(ri, ci)
            txt = row[ci] if ci < len(row) else ''
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
            set_cell_border(cell)
            if ri == 0:
                cell_shade(cell, HDR_BG)
                add_inline(p, txt, size=9.5, bold=True)
            else:
                bg = "FFFFFF" if ri % 2 == 1 else "F5F7FA"
                cell_shade(cell, bg)
                add_inline(p, txt, size=9.5)
    p_after = doc.add_paragraph(); p_after.paragraph_format.space_after = Pt(4)

def new_num_id(doc, start=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(a.get(qn('w:abstractNumId'))) for a in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    new_abs = (max(abstract_ids) + 1) if abstract_ids else 0
    new_num = (max(num_ids) + 1) if num_ids else 1
    abstractNum = OxmlElement('w:abstractNum'); abstractNum.set(qn('w:abstractNumId'), str(new_abs))
    mlt = OxmlElement('w:multiLevelType'); mlt.set(qn('w:val'), 'singleLevel'); abstractNum.append(mlt)
    lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), '0')
    s = OxmlElement('w:start'); s.set(qn('w:val'), str(start)); lvl.append(s)
    nf = OxmlElement('w:numFmt'); nf.set(qn('w:val'), 'decimal'); lvl.append(nf)
    lt = OxmlElement('w:lvlText'); lt.set(qn('w:val'), '%1.'); lvl.append(lt)
    jc = OxmlElement('w:lvlJc'); jc.set(qn('w:val'), 'left'); lvl.append(jc)
    pPr = OxmlElement('w:pPr'); ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480'); ind.set(qn('w:hanging'), '480')
    pPr.append(ind); lvl.append(pPr)
    abstractNum.append(lvl)
    first_num = numbering.find(qn('w:num'))
    if first_num is not None: first_num.addprevious(abstractNum)
    else: numbering.append(abstractNum)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(new_num))
    aref = OxmlElement('w:abstractNumId'); aref.set(qn('w:val'), str(new_abs)); num.append(aref)
    numbering.append(num)
    return new_num

def apply_numbering(paragraph, num_id, ilvl=0):
    pPr = paragraph._p.get_or_add_pPr()
    old = pPr.find(qn('w:numPr'))
    if old is not None: pPr.remove(old)
    numPr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'), str(ilvl)); numPr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'), str(num_id)); numPr.append(ni)
    pPr.append(numPr)

def add_page_number(doc):
    """在页脚中间加页码"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()
    set_font(run, size=9, color=MUTED)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'; instrText.set(qn('xml:space'), 'preserve')
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def add_header(doc, title_text):
    """在页眉右侧加系统名"""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.clear()
    r = p.add_run(title_text); set_font(r, size=9, color=MUTED)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), HAIR)
    pBdr.append(b); pPr.append(pBdr)

def severity_badge(doc, text, sev):
    """漏洞标题：带颜色左色条 + 等级角标"""
    ink_c, bg_c = SEVERITY_COLORS.get(sev, (INK, "F0F4FF"))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.35); shade(p, bg_c)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '36')
    lf.set(qn('w:space'), '6'); lf.set(qn('w:color'), ink_c)
    pBdr.append(lf); pPr.append(pBdr)
    r = p.add_run(text); set_font(r, size=14, bold=True, color=ink_c)
    badge = p.add_run(f"  [{sev}]"); set_font(badge, size=10, bold=True, color=ink_c)
    return p

def severity_from_heading(text):
    for sev in SEVERITY_COLORS:
        if sev in text:
            return sev
    return None

def main():
    doc = Document()
    normal = doc.styles['Normal']; normal.font.name = LATIN; normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.85); sec.right_margin = Cm(2.85)

    lines = open(SRC, encoding='utf-8').read().split('\n')

    # ── 解析封面 ────────────────────────────────────────────
    title = ""; meta = []; i = 0
    if lines and lines[0].lstrip().startswith('#'):
        title = lines[0].lstrip('#').strip(); i = 1
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith('>'):
            meta.append(s.lstrip('>').strip()); i += 1
        elif s == '':
            i += 1
        else:
            break
    while i < len(lines) and not lines[i].lstrip().startswith('## '):
        i += 1
    body_start = i

    add_page_number(doc)
    add_header(doc, title)

    # ── 封面 ─────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    # 大标题
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); set_font(r, size=28, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(8)

    # 细分隔线
    hr = doc.add_paragraph(); hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr.paragraph_format.space_after = Pt(16)
    pBdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), INK)
    pBdr.append(b); hr._p.get_or_add_pPr().append(pBdr)

    # meta 行（统计行单独用漏洞表渲染，此处跳过）
    stats_line = next((m for m in meta if '严重' in m or '高危' in m), None)
    for m in meta:
        if m is stats_line:
            continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, m, size=11, color=MUTED)
        p.paragraph_format.space_after = Pt(4)

    # 从正文 H2 标题提取漏洞列表
    vuln_rows = []
    for ln in lines[body_start:]:
        m2 = re.match(r'^##\s+(.*)$', ln.strip())
        if m2:
            text = strip_html(m2.group(1))
            sev = severity_from_heading(text)
            if sev:
                clean = re.sub(r'[（(](严重|高危|中危|低危|信息)[)）]', '', text).strip()
                vuln_rows.append((clean, sev))

    # 漏洞汇总表
    if vuln_rows:
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
        label_p = doc.add_paragraph()
        label_r = label_p.add_run("漏洞统计"); set_font(label_r, size=12, bold=True, color="1A1A1A")
        label_p.paragraph_format.space_after = Pt(6)

        tbl = doc.add_table(rows=1 + len(vuln_rows), cols=3)
        tbl.style = 'Table Grid'; tbl.autofit = True

        # 表头
        hdr_texts = ["序号", "漏洞名称", "风险等级"]
        for ci, txt in enumerate(hdr_texts):
            cell = tbl.cell(0, ci); cell.text = ''
            cell_shade(cell, HDR_BG); set_cell_border(cell)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
            r = p.add_run(txt); set_font(r, size=10, bold=True, color=INK)

        # 数据行
        for ri, (name, sev) in enumerate(vuln_rows, 1):
            ink_c, bg_c = SEVERITY_COLORS.get(sev, (INK, "FFFFFF"))
            row_bg = "FFFFFF" if ri % 2 == 1 else "F5F7FA"

            c0 = tbl.cell(ri, 0); c0.text = ''
            cell_shade(c0, row_bg); set_cell_border(c0)
            p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(2); p0.paragraph_format.space_before = Pt(2)
            r0 = p0.add_run(str(ri)); set_font(r0, size=10, color=MUTED); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

            c1 = tbl.cell(ri, 1); c1.text = ''
            cell_shade(c1, row_bg); set_cell_border(c1)
            p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(2); p1.paragraph_format.space_before = Pt(2)
            r1 = p1.add_run(name); set_font(r1, size=10, color="1A1A1A")

            c2 = tbl.cell(ri, 2); c2.text = ''
            cell_shade(c2, bg_c); set_cell_border(c2)
            p2 = c2.paragraphs[0]; p2.paragraph_format.space_after = Pt(2); p2.paragraph_format.space_before = Pt(2)
            r2 = p2.add_run(sev); set_font(r2, size=10, bold=True, color=ink_c); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 列宽：序号窄、名称宽、等级中
        from docx.oxml import OxmlElement as OE
        tbl.columns[0].width = Cm(1.2)
        tbl.columns[1].width = Cm(9.5)
        tbl.columns[2].width = Cm(2.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # ── 正文 ─────────────────────────────────────────────────
    i = body_start; in_code = False; code_buf = []; table_buf = []; current_num_id = None

    def flush_table_buf():
        nonlocal table_buf
        if table_buf: flush_table(doc, table_buf); table_buf = []

    while i < len(lines):
        ln = lines[i]

        # 代码块
        if ln.strip().startswith('```'):
            if not in_code:
                flush_table_buf(); in_code = True; code_buf = []
            else:
                add_code_block(doc, code_buf); in_code = False; code_buf = []
            current_num_id = None; i += 1; continue
        if in_code:
            if '```' in ln:
                idx = ln.index('```'); content_part = ln[:idx].rstrip()
                if content_part.strip(): code_buf.append(content_part)
                add_code_block(doc, code_buf); in_code = False; code_buf = []
                i += 1; continue
            code_buf.append(ln); i += 1; continue

        stripped = ln.strip()

        # 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buf.append(stripped); current_num_id = None; i += 1; continue
        else:
            flush_table_buf()

        if stripped == '': i += 1; continue

        # 有序列表
        mn = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if mn:
            if current_num_id is None: current_num_id = new_num_id(doc)
            p = doc.add_paragraph(); apply_numbering(p, current_num_id)
            p.paragraph_format.left_indent = Cm(0.74); p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, mn.group(2), size=10.5); i += 1; continue
        current_num_id = None

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1)); text = strip_html(m.group(2))
            sev = severity_from_heading(text)
            if level == 2 and sev:
                # 漏洞章节标题：带颜色色块
                clean = re.sub(r'[（(](严重|高危|中危|低危|信息)[)）]', '', text).strip()
                severity_badge(doc, clean, sev)
            elif level == 2:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
                r = p.add_run(text); set_font(r, size=15, bold=True, color=INK)
                left_bar(p, color=INK, sz='28')
            elif level == 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
                r = p.add_run(text); set_font(r, size=11, bold=True, color=ACCENT)
                left_bar(p, color=ACCENT, sz='16', space='6')
            else:
                p = doc.add_paragraph()
                r = p.add_run(text); set_font(r, size=10.5, bold=True, color=INK)
            i += 1; continue

        # 分割线
        if re.match(r'^---+$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
            pBdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '1'); b.set(qn('w:color'), HAIR)
            pBdr.append(b); p._p.get_or_add_pPr().append(pBdr)
            i += 1; continue

        # 引用块
        if stripped.startswith('>'):
            txt = stripped.lstrip('>').strip()
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            pBdr = OxmlElement('w:pBdr'); lf = OxmlElement('w:left')
            lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '16')
            lf.set(qn('w:space'), '8'); lf.set(qn('w:color'), ACCENT)
            pBdr.append(lf); p._p.get_or_add_pPr().append(pBdr)
            shade(p, 'F0F4FF'); add_inline(p, txt, size=10)
            i += 1; continue

        # 无序列表
        mb = re.match(r'^[-*]\s+(.*)$', stripped)
        if mb:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, mb.group(1), size=10.5); i += 1; continue

        # 普通段落
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        add_inline(p, stripped, size=10.5); i += 1

    flush_table_buf()
    os.makedirs(os.path.dirname(OUT) or '.', exist_ok=True)
    doc.save(OUT)
    print("SAVED:", OUT, "size=", os.path.getsize(OUT))

if __name__ == '__main__':
    main()
