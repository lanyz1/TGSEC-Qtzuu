# -*- coding: utf-8 -*-
"""安全评估报告渲染器：Markdown → HTML + DOCX（模块 + CLI）。

自包含渲染器，两种格式共用同一 parse_markdown 块模型：
- render_html：内联 CSS、无外链、打印友好、危害等级配色徽章；
- render_docx：python-docx，封面 + 危害配色 + 表格（docx 底层工具借鉴 Pentest-Lyan render_docx.py）。

Markdown 约定：首行 `# 标题` 为封面标题，紧随的 `> ...` 引用行为封面元信息（目标/时间/统计），
之后 `---` 进入正文。正文支持标题(H1-H6)/段落/表格/围栏代码/有序·无序列表/引用/分割线/
内联(加粗 ** / 代码 ` / 链接 []() / 删除 ~~)。危害徽章按标题中的 严重/高危/中危/低危/信息 着色。

CLI：
    python render_report.py <input.md> [--html out.html] [--docx out.docx]
    （缺 --html/--docx 时默认与输入同名输出两者）
"""

import argparse
import html as _html
import os
import re
import sys

# ---- 危害等级配色（中文等级 → (前景, 背景)），与 critical/high/medium/low/info 对齐 ----
SEVERITY_COLORS = {
    "严重": ("C0392B", "FDECEA"),
    "高危": ("D35400", "FEF0E7"),
    "中危": ("B7770D", "FEF9E7"),
    "低危": ("1A7A4A", "E9F7EF"),
    "信息": ("2471A3", "EBF5FB"),
}
SEVERITY_ORDER = ["严重", "高危", "中危", "低危", "信息"]

INK = "1F3A68"      # 深蓝：正文标题
ACCENT = "2E6FD4"   # 亮蓝：次级标题
MUTED = "595959"    # 灰：元信息
HDR_BG = "D9E2F3"   # 表头浅蓝
CODE_BG = "F4F4F4"  # 代码块底
HAIR = "BFBFBF"     # 分割线

_SEV_RE = re.compile(r"(严重|高危|中危|低危|信息)")


def severity_in(text):
    m = _SEV_RE.search(text or "")
    return m.group(1) if m else None


# ============================ Markdown 解析 ============================

def parse_cover(lines):
    """解析封面：首个 `# 标题` + 紧随的 `> 元信息` 行。返回 (title, meta_lines, body_start)。"""
    title, meta, i = "", [], 0
    n = len(lines)
    while i < n and lines[i].strip() == "":
        i += 1
    if i < n and lines[i].lstrip().startswith("# ") and not lines[i].lstrip().startswith("##"):
        title = lines[i].lstrip("#").strip()
        i += 1
    while i < n:
        s = lines[i].strip()
        if s.startswith(">"):
            meta.append(s.lstrip(">").strip())
            i += 1
        elif s == "":
            i += 1
        else:
            break
    # 跳过封面后到首个正文块（越过 --- 分隔线）
    while i < n and (lines[i].strip() == "" or re.match(r"^-{3,}$", lines[i].strip())):
        i += 1
    return title, meta, i


def parse_markdown(lines, start=0):
    """把正文行解析为块序列。每块为 dict，type ∈
    heading/para/table/code/list/quote/hr。"""
    blocks = []
    i, n = start, len(lines)
    para_buf = []

    def flush_para():
        if para_buf:
            blocks.append({"type": "para", "text": " ".join(para_buf).strip()})
            para_buf.clear()

    while i < n:
        raw = lines[i]
        s = raw.strip()

        # 围栏代码
        if s.startswith("```"):
            flush_para()
            lang = s[3:].strip()
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            blocks.append({"type": "code", "lines": code, "lang": lang})
            continue

        # 表格（连续 | 行）
        if s.startswith("|") and s.endswith("|"):
            flush_para()
            tbl = []
            while i < n and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            rows = []
            for r in tbl:
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                if all(re.match(r"^:?-+:?$", c or "") for c in cells):
                    continue  # 分隔行
                rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # 分割线
        if re.match(r"^(-{3,}|\*{3,})$", s):
            flush_para()
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            flush_para()
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue

        # 引用
        if s.startswith(">"):
            flush_para()
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "quote", "text": " ".join(quote).strip()})
            continue

        # 有序列表
        if re.match(r"^\d+\.\s+", s):
            flush_para()
            items = []
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append({"type": "list", "ordered": True, "items": items})
            continue

        # 无序列表
        if re.match(r"^[-*]\s+", s):
            flush_para()
            items = []
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append({"type": "list", "ordered": False, "items": items})
            continue

        # 空行 → 段落分隔
        if s == "":
            flush_para()
            i += 1
            continue

        # 普通段落行（累积）
        para_buf.append(s)
        i += 1

    flush_para()
    return blocks


# ============================ HTML 渲染 ============================

_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|~~[^~]+~~)")
_HTML_TAG_RE = re.compile(r"<[^>]+>")


def _inline_html(text):
    """行内 markdown → HTML（先转义再套标签），支持 **加粗** `代码` [文本](url) ~~删除~~。"""
    text = _HTML_TAG_RE.sub("", text or "")
    out, pos = [], 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append(_html.escape(text[pos:m.start()]))
        tok = m.group(0)
        if tok.startswith("**"):
            out.append("<strong>%s</strong>" % _html.escape(tok[2:-2]))
        elif tok.startswith("`"):
            out.append("<code>%s</code>" % _html.escape(tok[1:-1]))
        elif tok.startswith("~~"):
            out.append("<del>%s</del>" % _html.escape(tok[2:-2]))
        elif tok.startswith("["):
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            if mm:
                out.append('<a href="%s">%s</a>' % (_html.escape(mm.group(2)), _html.escape(mm.group(1))))
        pos = m.end()
    if pos < len(text):
        out.append(_html.escape(text[pos:]))
    return "".join(out)


def _slug(text, used):
    base = re.sub(r"[^\w一-鿿]+", "-", (text or "").strip()).strip("-").lower() or "sec"
    s, k = base, 1
    while s in used:
        k += 1
        s = "%s-%d" % (base, k)
    used.add(s)
    return s


_HTML_CSS = """
:root{--ink:#1F3A68;--accent:#2E6FD4;--muted:#595959;--hair:#D0D7E2;--code-bg:#F4F6F9;--bg:#ffffff;--fg:#1a1a1a;}
*{box-sizing:border-box;}
body{margin:0;background:#eef1f6;color:var(--fg);font-family:"Segoe UI","Microsoft YaHei","微软雅黑",system-ui,sans-serif;font-size:15px;line-height:1.7;}
.page{max-width:960px;margin:24px auto;background:var(--bg);padding:48px 56px;box-shadow:0 2px 16px rgba(0,0,0,.08);border-radius:6px;}
.cover{text-align:center;padding:24px 0 20px;border-bottom:3px solid var(--ink);margin-bottom:8px;}
.cover h1{color:var(--ink);font-size:30px;margin:8px 0 14px;letter-spacing:1px;}
.cover .meta{color:var(--muted);font-size:14px;margin:3px 0;}
.cover .stats{margin-top:16px;display:flex;gap:8px;justify-content:center;flex-wrap:wrap;}
.badge{display:inline-block;padding:3px 12px;border-radius:14px;font-weight:600;font-size:13px;}
.toc{background:#f7f9fc;border:1px solid var(--hair);border-radius:6px;padding:14px 20px;margin:22px 0;}
.toc-title{font-weight:700;color:var(--ink);margin-bottom:6px;}
.toc a{color:var(--accent);text-decoration:none;margin-right:4px;}
.toc a:hover{text-decoration:underline;}
h2{color:var(--ink);font-size:21px;margin:32px 0 12px;padding-left:12px;border-left:5px solid var(--ink);}
h3{font-size:17px;margin:22px 0 8px;padding:8px 12px;border-left:5px solid var(--accent);background:#f0f4ff;border-radius:0 4px 4px 0;}
h3.vuln{border-left-width:6px;}
h4{color:var(--ink);font-size:15px;margin:16px 0 6px;}
h5,h6{color:var(--muted);font-size:14px;margin:12px 0 4px;}
p{margin:8px 0;}
code{background:var(--code-bg);color:#C0392B;font-family:Consolas,Monaco,monospace;font-size:.9em;padding:1px 5px;border-radius:3px;}
pre{background:var(--code-bg);border-left:4px solid var(--accent);border-radius:0 4px 4px 0;padding:12px 14px;overflow-x:auto;margin:10px 0;}
pre code{background:none;color:#2C3E50;padding:0;font-size:13px;line-height:1.5;}
.table-wrap{overflow-x:auto;margin:12px 0;}
table{border-collapse:collapse;width:100%;font-size:14px;}
th,td{border:1px solid var(--hair);padding:7px 10px;text-align:left;vertical-align:top;}
th{background:#D9E2F3;color:var(--ink);font-weight:700;}
tbody tr:nth-child(even){background:#F5F7FA;}
blockquote{margin:10px 0;padding:8px 14px;border-left:4px solid var(--accent);background:#f0f4ff;color:#33415c;border-radius:0 4px 4px 0;}
ul,ol{margin:8px 0;padding-left:26px;}
li{margin:3px 0;}
hr{border:none;border-top:1px solid var(--hair);margin:18px 0;}
a{color:var(--accent);}
@media print{body{background:#fff;}.page{box-shadow:none;margin:0;max-width:none;padding:0;}}
"""


def render_html(blocks, title, meta_lines):
    used = set()
    parts = ['<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">',
             '<meta name="viewport" content="width=device-width,initial-scale=1">',
             "<title>%s</title>" % _html.escape(title or "安全评估报告"),
             "<style>%s</style></head><body><div class=\"page\">" % _HTML_CSS]

    # 封面
    parts.append('<div class="cover">')
    parts.append("<h1>%s</h1>" % _html.escape(title or "安全评估报告"))
    stats_line = None
    for m in meta_lines:
        if "漏洞统计" in m or ("严重" in m and "高危" in m):
            stats_line = m
            continue
        parts.append('<div class="meta">%s</div>' % _inline_html(m))
    if stats_line:
        parts.append('<div class="stats">')
        label = stats_line.split(":", 1)[-1] if ":" in stats_line else stats_line.split("：", 1)[-1]
        for seg in re.split(r"[/、]", label):
            sev = severity_in(seg)
            if sev:
                fg, bg = SEVERITY_COLORS[sev]
                parts.append('<span class="badge" style="color:#%s;background:#%s;">%s</span>'
                             % (fg, bg, _html.escape(seg.strip())))
        parts.append("</div>")
    parts.append("</div>")

    # 目录（H2）——预分配 slug，TOC 与正文锚点复用同一映射（顺序一致）
    h2_texts = [b["text"] for b in blocks if b["type"] == "heading" and b["level"] == 2]
    h2_slugs = [_slug(t, used) for t in h2_texts]
    if h2_texts:
        parts.append('<div class="toc"><div class="toc-title">目录</div><div>')
        for t, sl in zip(h2_texts, h2_slugs):
            parts.append('<a href="#%s">%s</a>' % (sl, _html.escape(t)))
        parts.append("</div></div>")
    h2_idx = 0

    # 正文
    for b in blocks:
        t = b["type"]
        if t == "heading":
            lvl = min(max(b["level"], 1), 6)
            text = b["text"]
            sev = severity_in(text) if lvl == 3 else None
            anchor = ""
            if lvl == 2:
                anchor = ' id="%s"' % h2_slugs[h2_idx]
                h2_idx += 1
            if sev and lvl == 3:
                fg, bg = SEVERITY_COLORS[sev]
                parts.append('<h3 class="vuln" style="border-left-color:#%s;background:#%s;">%s</h3>'
                             % (fg, bg, _inline_html(text)))
            else:
                parts.append("<h%d%s>%s</h%d>" % (lvl, anchor, _inline_html(text), lvl))
        elif t == "para":
            parts.append("<p>%s</p>" % _inline_html(b["text"]))
        elif t == "quote":
            parts.append("<blockquote>%s</blockquote>" % _inline_html(b["text"]))
        elif t == "hr":
            parts.append("<hr>")
        elif t == "code":
            parts.append("<pre><code>%s</code></pre>"
                         % _html.escape("\n".join(b["lines"])))
        elif t == "list":
            tag = "ol" if b["ordered"] else "ul"
            parts.append("<%s>" % tag)
            for it in b["items"]:
                parts.append("<li>%s</li>" % _inline_html(it))
            parts.append("</%s>" % tag)
        elif t == "table":
            parts.append('<div class="table-wrap"><table>')
            for ri, row in enumerate(b["rows"]):
                if ri == 0:
                    parts.append("<thead><tr>")
                    for c in row:
                        parts.append("<th>%s</th>" % _inline_html(c))
                    parts.append("</tr></thead><tbody>")
                else:
                    parts.append("<tr>")
                    for c in row:
                        parts.append("<td>%s</td>" % _inline_html(c))
                    parts.append("</tr>")
            parts.append("</tbody></table></div>")
    parts.append("</div></body></html>")
    return "".join(parts)


# ============================ DOCX 渲染 ============================

def render_docx(blocks, title, meta_lines, out_path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    CJK, LATIN, MONO = "微软雅黑", "Segoe UI", "Consolas"

    def set_font(run, name=LATIN, cjk=CJK, size=None, bold=None, color=None):
        run.font.name = name
        r = run._element.get_or_add_rPr()
        rFonts = r.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            r.append(rFonts)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), cjk)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)

    def shade(paragraph, fill):
        pPr = paragraph._p.get_or_add_pPr()
        old = pPr.find(qn("w:shd"))
        if old is not None:
            pPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def cell_shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:shd"))
        if old is not None:
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def left_bar(paragraph, color=INK, sz="24", space="8"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        lf = OxmlElement("w:left")
        lf.set(qn("w:val"), "single")
        lf.set(qn("w:sz"), sz)
        lf.set(qn("w:space"), space)
        lf.set(qn("w:color"), color)
        pBdr.append(lf)
        pPr.append(pBdr)

    def set_cell_border(cell, color="D0D7E2", sz="4"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBdr = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement("w:%s" % side)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tcBdr.append(el)
        tcPr.append(tcBdr)

    HTML_TAG_RE = re.compile(r"<[^>]+>")
    INLINE_RE = _INLINE_RE

    def add_inline(paragraph, text, size=10.5, color=None, bold=False):
        text = HTML_TAG_RE.sub("", text or "")
        pos = 0
        for m in INLINE_RE.finditer(text):
            if m.start() > pos:
                r = paragraph.add_run(text[pos:m.start()])
                set_font(r, size=size, color=color, bold=bold if bold else None)
            tok = m.group(0)
            if tok.startswith("**"):
                r = paragraph.add_run(tok[2:-2])
                set_font(r, size=size, bold=True, color=color)
            elif tok.startswith("`"):
                r = paragraph.add_run(tok[1:-1])
                set_font(r, name=MONO, cjk=MONO, size=size - 0.5, color="C0392B")
            elif tok.startswith("~~"):
                r = paragraph.add_run(tok[2:-2])
                set_font(r, size=size, color=color)
                r.font.strike = True
            elif tok.startswith("["):
                mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
                if mm:
                    r = paragraph.add_run(mm.group(1))
                    set_font(r, size=size, color="2471A3")
            pos = m.end()
        if pos < len(text):
            r = paragraph.add_run(text[pos:])
            set_font(r, size=size, color=color, bold=bold if bold else None)

    def add_code_block(doc, lines):
        for ln in lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.4)
            shade(p, CODE_BG)
            left_bar(p, color=ACCENT, sz="16", space="6")
            r = p.add_run(ln if ln != "" else " ")
            set_font(r, name=MONO, cjk=MONO, size=9, color="2C3E50")

    def add_table(doc, rows):
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell = tbl.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                set_cell_border(cell)
                txt = row[ci] if ci < len(row) else ""
                if ri == 0:
                    cell_shade(cell, HDR_BG)
                    add_inline(p, txt, size=9.5, bold=True)
                else:
                    cell_shade(cell, "FFFFFF" if ri % 2 == 1 else "F5F7FA")
                    add_inline(p, txt, size=9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def severity_badge(doc, text, sev):
        fg, bg = SEVERITY_COLORS.get(sev, (INK, "F0F4FF"))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.35)
        shade(p, bg)
        left_bar(p, color=fg, sz="36", space="6")
        r = p.add_run(re.sub(r"[（(](严重|高危|中危|低危|信息)[)）]", "", HTML_TAG_RE.sub("", text)).strip())
        set_font(r, size=13, bold=True, color=fg)
        badge = p.add_run("  [%s]" % sev)
        set_font(badge, size=10, bold=True, color=fg)

    def add_page_number(doc):
        p = doc.sections[0].footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        set_font(run, size=9, color=MUTED)
        f1 = OxmlElement("w:fldChar")
        f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText")
        it.text = "PAGE"
        it.set(qn("xml:space"), "preserve")
        f2 = OxmlElement("w:fldChar")
        f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1)
        run._r.append(it)
        run._r.append(f2)

    def add_header(doc, text):
        p = doc.sections[0].header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        set_font(r, size=9, color=MUTED)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.4)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.6)

    add_page_number(doc)
    add_header(doc, title or "安全评估报告")

    # 封面
    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title or "安全评估报告")
    set_font(r, size=26, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(8)
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr.paragraph_format.space_after = Pt(14)
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), INK)
    pBdr.append(b)
    hr._p.get_or_add_pPr().append(pBdr)
    for m in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, m, size=11, color=MUTED)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # 正文
    for blk in blocks:
        t = blk["type"]
        if t == "heading":
            lvl, text = blk["level"], blk["text"]
            sev = severity_in(text) if lvl == 3 else None
            if lvl == 3 and sev:
                severity_badge(doc, text, sev)
            elif lvl <= 2:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(HTML_TAG_RE.sub("", text))
                set_font(r, size=15 if lvl == 2 else 17, bold=True, color=INK)
                left_bar(p, color=INK, sz="28")
            elif lvl == 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(HTML_TAG_RE.sub("", text))
                set_font(r, size=12, bold=True, color=ACCENT)
                left_bar(p, color=ACCENT, sz="16", space="6")
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(HTML_TAG_RE.sub("", text))
                set_font(r, size=11, bold=True, color=INK)
        elif t == "para":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, blk["text"], size=10.5)
        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(3)
            shade(p, "F0F4FF")
            left_bar(p, color=ACCENT, sz="16", space="8")
            add_inline(p, blk["text"], size=10)
        elif t == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBdr = OxmlElement("w:pBdr")
            bb = OxmlElement("w:bottom")
            bb.set(qn("w:val"), "single")
            bb.set(qn("w:sz"), "4")
            bb.set(qn("w:space"), "1")
            bb.set(qn("w:color"), HAIR)
            pBdr.append(bb)
            p._p.get_or_add_pPr().append(pBdr)
        elif t == "code":
            add_code_block(doc, blk["lines"])
        elif t == "list":
            style = "List Number" if blk["ordered"] else "List Bullet"
            for it in blk["items"]:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(2)
                add_inline(p, it, size=10.5)
        elif t == "table":
            add_table(doc, blk["rows"])

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
    doc.save(out_path)


# ============================ 渲染入口 ============================

def render_file(md_text, html_out=None, docx_out=None):
    """解析 md 文本并按需渲染。返回 (title, 已产出路径列表)。"""
    lines = md_text.split("\n")
    title, meta_lines, body_start = parse_cover(lines)
    blocks = parse_markdown(lines, body_start)
    produced = []
    if html_out:
        with open(html_out, "w", encoding="utf-8") as f:
            f.write(render_html(blocks, title, meta_lines))
        produced.append(html_out)
    if docx_out:
        render_docx(blocks, title, meta_lines, docx_out)
        produced.append(docx_out)
    return title, produced


def main():
    p = argparse.ArgumentParser(description="安全评估报告渲染器 Markdown → HTML + DOCX")
    p.add_argument("input", help="输入 Markdown 报告 .md")
    p.add_argument("--html", default=None, help="HTML 输出路径（缺省与输入同名 .html）")
    p.add_argument("--docx", default=None, help="DOCX 输出路径（缺省与输入同名 .docx）")
    args = p.parse_args()
    if not args.input.endswith(".md"):
        sys.exit("输入必须是 .md 文件")
    html_out = args.html or (args.input[:-3] + ".html")
    docx_out = args.docx or (args.input[:-3] + ".docx")
    with open(args.input, encoding="utf-8") as f:
        md = f.read()
    title, produced = render_file(md, html_out, docx_out)
    for pth in produced:
        print("SAVED: %s (%d bytes)" % (pth, os.path.getsize(pth)))


if __name__ == "__main__":
    main()
